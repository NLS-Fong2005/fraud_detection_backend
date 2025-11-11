# <--- Imports --->
import os
import shutil
import time
import pandas as pd
from docx import Document
from typing import (
    List
)
from weaviate.util import generate_uuid5

from src.machine_learning.llm_rag_method.vector_store.vector_database import vector_collection

class DocumentLoader:
    def __init__(self):
        self.GUIDELINES_DIRECTORY: str = os.path.abspath("./guidelines")
        self.TRAINED_DIRECTORY: str = os.path.abspath("./guidelines/trained_data")

    def __get_processing_time__(self, start_time: float, end_time: float) -> str:
        time_taken = (end_time - start_time)
        if (time_taken < 0):
            time_taken *= -1
        
        if time_taken > 60:
            minutes: float = (time_taken) / 60 #type: ignore
            seconds: float = (time_taken) % 60 #type: ignore
            result: str = f"{int(minutes)}m {seconds}s"
        else:
            result: str = f"{time_taken}s"
        return result

    def __retrieve_files__(self) -> List[str]:
        return [file for file in os.listdir(path=self.GUIDELINES_DIRECTORY) if file.endswith(".docx")]
        
    def __extract_tables_from_document__(self, document) -> dict:
        table_headers = []
        row_texts = {}
        extracted_tables = {}

        for table_number, table in enumerate(document.tables):
            row_texts[table_number] = []
            for row_index, row in enumerate(table.rows):
                if row_index == 0:
                    table_headers.append([cell.text for cell in row.cells])
                else:
                    row = [cell.text for cell in row.cells]
                    row_texts[table_number].append(row) 

        for table_number, headings in enumerate(table_headers):
            extracted_tables[table_number] = {}
            
            for col_number, header in enumerate(headings):
                extracted_tables[table_number].update(
                    {header : [row_text[col_number] for row_text in row_texts[table_number]]}
                )
        return extracted_tables

    def __extract_text_from_document__(self, document_path: str) -> dict:
        document = Document(document_path)
        if (len(document.paragraphs) == 0):
            raise Exception("Document is empty! No text to extract.")

        result = {
            "Paragraphs": "\n".join([paragraph.text for paragraph in document.paragraphs]),
            "Tables": self.__extract_tables_from_document__(document)
        }
        return result
        
    def add_documents(self):
        def document_exists(collection, doc_uuid: str) -> bool:
            return collection.data.exists(doc_uuid)

        collection = vector_collection.get_vector_collection()
        documents = self.__retrieve_files__()

        print(f"\n--- Found {len(documents)} to process ---\n")
        for number, file_name in enumerate(documents):
            print(f"{number + 1}: {file_name}")

        total_start_time: float = time.time()
        for document in documents: 
            header = document.rstrip(f".docx")

            tables_from_docx = {}
            for table_number, table in self.__extract_text_from_document__(f"{self.GUIDELINES_DIRECTORY}/{document}")["Tables"].items():
                df = pd.DataFrame.from_dict(table)
                tables_from_docx[table_number] = df

            extracted_document = f"""
                # Extracted Paragraphs
                {"-" * 30}
                {self.__extract_text_from_document__(f"{self.GUIDELINES_DIRECTORY}/{document}")["Paragraphs"]}
                {"-" * 30}
                
                # Extracted Tables
                {"-" * 30}
                {tables_from_docx}
                {"-" * 30}
            """

            print(f"\n--- Processing Guideline Document: {header} ---\n")
            guideline_object = {
                "header": header,
                "info": extracted_document,
            }
            knowledge_uuid = generate_uuid5(guideline_object) 

            if not document_exists(collection, knowledge_uuid):
                collection.data.insert(
                    properties=guideline_object,
                    uuid=knowledge_uuid
                )
                print(f"Document has been added under {knowledge_uuid}")
                destination_path = os.path.join(self.TRAINED_DIRECTORY, os.path.basename(f"{self.GUIDELINES_DIRECTORY}/{document}"))
                shutil.move(f"{self.GUIDELINES_DIRECTORY}/{document}", destination_path)
                print(f"File '{document}' moved successfully to '{self.TRAINED_DIRECTORY}")
            else:
                print(f"Document already exists under {knowledge_uuid}!")
        total_end_time = time.time()
        print(f"Time taken for processing {len(documents)} number of Documents: {self.__get_processing_time__(start_time=total_start_time, end_time=total_end_time)}")

document_loader = DocumentLoader()

if __name__ == "__main__":
    document_loader.add_documents()